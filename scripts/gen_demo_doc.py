# -*- coding: utf-8 -*-
"""
毕数飞行平台 - 全流程系统演示汇报流程与讲解话术 Word 文档生成脚本
生成目标: 设计文档/毕数飞行平台_系统演示汇报流程与讲解话术.docx
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(22)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        run2.font.name = '微软雅黑'
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run2.font.size = Pt(10.5)
        run2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

def add_h1(doc, text):
    h = doc.add_heading(level=1)
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after = Pt(6)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)
    return h

def add_h2(doc, text):
    h = doc.add_heading(level=2)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(4)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    return h

def add_h3(doc, text):
    h = doc.add_heading(level=3)
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(2)
    run = h.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return h

def add_p(doc, text, bold=False, italic=False, color=None, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(10)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    return p

def add_callout(doc, title, content_list, border_color='1A56C4', bg_color='F0F4FA'):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell.width = Cm(16.5)
    set_cell_shading(cell, bg_color)
    set_cell_margins(cell, top=140, bottom=140, left=220, right=220)

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r_title = p.add_run(f"【{title}】\n")
    r_title.font.name = '微软雅黑'
    r_title.bold = True
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(0x0F, 0x4C, 0x81)

    for item in content_list:
        p_item = cell.add_paragraph()
        p_item.paragraph_format.space_after = Pt(2)
        p_item.paragraph_format.line_spacing = 1.15
        
        if isinstance(item, tuple):
            prefix, txt = item
            r_pre = p_item.add_run(prefix)
            r_pre.font.name = '微软雅黑'
            r_pre.bold = True
            r_pre.font.size = Pt(9.5)
            r_pre.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            
            r_txt = p_item.add_run(txt)
            r_txt.font.name = '微软雅黑'
            r_txt.font.size = Pt(9.5)
            r_txt.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r_txt = p_item.add_run(item)
            r_txt.font.name = '微软雅黑'
            r_txt.font.size = Pt(9.5)
            r_txt.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def add_table_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '1A56C4')
        set_cell_margins(cell, 120, 120, 160, 160)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = '微软雅黑'
            r.font.size = Pt(9.5)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for row_idx, r_data in enumerate(rows):
        for col_idx, text in enumerate(r_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(text)
            set_cell_margins(cell, 100, 100, 140, 140)
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F8FAFC')
            p = cell.paragraphs[0]
            for r in p.runs:
                r.font.name = '微软雅黑'
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return table

def generate_demo_word_doc(out_path):
    doc = Document()
    
    # 页面边距设置 (A4, 2.2cm 边距)
    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.2)

    # 1. 标题与副标题
    add_title(doc, "毕数飞行平台 · 全流程系统演示汇报流程与讲解话术", 
              "编制版本：V1.0  |  编制日期：2026年8月  |  用途：领导层决策、投资人路演、合作伙伴签约演示")

    # 2. 演示准备与环境配置
    add_h1(doc, "一、 演示准备与环境配置")
    add_p(doc, "为保障演示效果流畅、多端联动无缝衔接，建议在浏览器中预先打开以下 5 个系统标签页（推荐按 F11 开启全屏模式）：")

    env_headers = ["序号", "系统端名称", "文件/访问路径", "角色定位", "核心演示功能点"]
    env_rows = [
        ["1", "演示总控门户", "演示系统/index.html", "汇报主控入口", "平台定位、全景架构说明、4大核心演示路线一键直达"],
        ["2", "C端小程序", "演示系统/c-app.html", "学员 / 飞手 / 需求方", "0元免费培训、CAAC微信分账收银台、抢单履约、UOM合规、IM聊天"],
        ["3", "B端培训机构后台", "演示系统/b-training.html", "定点航校 / 考证机构", "防挂机微表情抽检、合格证明防伪生成、CAAC执照回传、人社补贴批量申报"],
        ["4", "B端用机企业后台", "演示系统/b-enterprise.html", "企业客户 / 需求方", "50km高压巡检发布、机组飞手调度、区块链合同存证调阅、对公结算"],
        ["5", "运营风控中枢", "演示系统/b-admin.html", "Super Admin / 监管方", "GIS全域电子沙盘、防跳单仲裁法庭、维修商准入、微信分账清算台账"]
    ]
    add_table_data(doc, env_headers, env_rows, [1.2, 2.8, 3.8, 3.2, 5.5])

    # 3. 总体演示顺序架构
    add_h1(doc, "二、 演示整体时序架构（黄金四幕法）")
    add_p(doc, "整体演示时长建议控制在 10~15 分钟，采用“起-承-转-合”的黄金叙事节奏，先讲飞手引流与变现造血，再讲企业发单合规履约，后亮出防跳单护城河，最后以大屏与清算台账收尾。")

    arch_headers = ["阶段", "幕次名称", "建议时长", "主演示系统", "核心说服点 / 商业逻辑"]
    arch_rows = [
        ["开篇", "平台全景与架构介绍", "1 分钟", "index.html", "痛点引入：低空经济爆发期，解决飞手合规难、发单撮合难、私下交易多的痛点。"],
        ["第一幕", "飞手成长与商业造血闭环", "3~4 分钟", "c-app ➔ b-training ➔ b-admin", "人社免费培训精准引流 ➔ 800元券转化CAAC考证 ➔ 微信分账20%佣金 ➔ 执照回传解锁接单。"],
        ["第二幕", "企业用机与全流程合规履约", "3~4 分钟", "b-enterprise ➔ c-app", "企业发布大型巡检 ➔ 飞手抢单资质拦截 ➔ 电子合同存证 ➔ UOM合规与GPS作业打卡。"],
        ["第三幕", "五维防跳单与风控法庭", "2~3 分钟", "c-app ➔ b-admin", "技术护城河：IM敏感词毫秒拦截 ➔ AXB隐私号录音AI质检 ➔ 仲裁扣除¥1000保证金。"],
        ["第四幕", "低空全域大屏与微信分账", "2~3 分钟", "b-admin", "全域GIS沙盘态势掌控 ➔ 维修商准入 ➔ ¥245万微信分账冻结资金池（零银行、防二清）。"],
        ["结尾", "总结与互动问答", "1~2 分钟", "index.html", "全链路已完成原型验证与前端工程化，随时可启动后端联调商业化落地。"]
    ]
    add_table_data(doc, arch_headers, arch_rows, [1.5, 3.2, 2.0, 3.8, 6.0])

    # 4. 逐幕详细演说词与操作指南
    add_h1(doc, "三、 逐幕详细演说词与操作指南")

    # 第一幕
    add_h2(doc, "第一幕：飞手成长与商业造血闭环（人社培训 ➔ CAAC考证分账 ➔ 补贴申报）")
    
    add_callout(doc, "步骤 1.1：C端小程序 · 零费用技能培训报名", [
        ("【演示操作】", "切换至 c-app.html ➔ 底部点击「学习中心」➔ 保持在「技能培训」Tab ➔ 点击「农业植保无人机飞防实训班」➔ 点击「免费报名」➔ 弹出身份核验。"),
        ("【界面呈现】", "展示人社定点机构、补贴政策说明、身份证与就业状态表单，提交后提示“报名成功，人社补贴资格已核验”。"),
        ("【标准讲解话术】", "“各位领导、专家好！低空经济的核心支撑是高素质持证飞手。毕数飞行平台的第一道造血引流入口，是对接政府人社部门的技能培训补贴政策。对于广大零基础求职者，平台提供经人社备案的定点培训班，学员在线填写身份信息即可‘全程0元免费报名’，培训费用由政府财政直补定点机构，从而为平台低成本聚集海量飞手私域流量。”")
    ])

    add_callout(doc, "步骤 1.2：培训机构端 · 防挂机微表情抽检与合格证明防伪生成", [
        ("【演示操作】", "切换至 b-training.html ➔ 点击左侧「班级与考勤」➔ 查看防挂机监控记录 ➔ 点击「生成合格证明」弹窗。"),
        ("【界面呈现】", "展示“线上防挂机微表情抓拍抽检 + 线下每日指纹真机打卡”双轨核验，弹窗展示带防伪二维码和 PT2026 编号的正式合格证书。"),
        ("【标准讲解话术】", "“培训机构在管理端严格落实人社监管要求。学员线上学习通过微表情抓拍防挂机，线下实操真机打卡。结业考核通过后，系统自动生成带唯一防伪二维码的培训合格证明。人社部门扫码即可穿透核验培训档案，这也是机构申请财政补贴的核心凭证。”")
    ])

    add_callout(doc, "步骤 1.3：C端高客单转化 · 800元CAAC考证券与微信分账收银台", [
        ("【演示操作】", "切换回 c-app.html ➔ 弹出「学员专属 800 元 CAAC 考证抵用券」➔ 点击「立即使用」进入报名收银台 ➔ 点击「微信支付 ¥4,200」。"),
        ("【界面呈现】", "收银台清晰注明：资金通道为‘微信支付分账资金托管’，结算方式为‘报名即分账 · 机构 T+1 结算’，包含按课时阶梯退费保障。"),
        ("【标准讲解话术】", "“学员在平台结业后，系统自动触发‘成长赋能弹窗’，发放专属 800 元抵用券，激励学员进一步考取民航局 CAAC 执照，实现高客单业务转化。学员付款后，资金进入微信支付分账专户冻结，平台直接获得 20% 服务费（¥1,600），机构获得 80% 培训款，实现‘报名即分账’的高效资金周转。”")
    ])

    add_callout(doc, "步骤 1.4：培训机构端 · CAAC执照回传与一键补贴申报", [
        ("【演示操作】", "切换至 b-training.html ➔ 点击「CAAC排考中心」➔ 找到学员‘张强’点击「回传执照」➔ 点击「补贴申报中心」展示打包申报。"),
        ("【界面呈现】", "回传执照后即时弹窗提示“已成功激活飞手接单资格”；补贴申报中心一键打包该批次执照信息向人社局申报。"),
        ("【标准讲解话术】", "“学员通过民航局考试后，机构在后台回传 CAAC 电子执照，平台即时为该学员升级为持证飞手并解锁抢单权限；同时机构在后台一键打包向人社部门申报证书补贴，形成完整的培训-考证-就业商业闭环。”")
    ])

    # 第二幕
    add_h2(doc, "第二幕：企业大型用机与全流程合规履约（发单 ➔ 抢单 ➔ 合同 ➔ UOM ➔ 验收）")

    add_callout(doc, "步骤 2.1：企业端发布大型飞行项目与机组调度", [
        ("【演示操作】", "切换至 b-enterprise.html ➔ 点击「飞行项目监控」➔ 点击「发布大型飞行项目」➔ 点击「机组飞手调度」查看指派。"),
        ("【界面呈现】", "展示德阳 50km 高压输电线路精细化巡检项目（标的 ¥38,000），机组成员执照与设备状态一目了然。"),
        ("【标准讲解话术】", "“第二板块是企业用机需求撮合。大型用机单位（如电网、测绘局、农业企业）在企业端发布复杂作业项目，指派机组人员，项目定金与尾款通过微信商户专户全额托管，保障资金安全。”")
    ])

    add_callout(doc, "步骤 2.2：C端飞手抢单 · 资质拦截与区块链电子合同签署", [
        ("【演示操作】", "切换至 c-app.html ➔ 点击「接单广场」➔ 找到该电力巡检订单 ➔ 点击「立即抢单」➔ 弹出「签署电子合同」➔ 签名确认。"),
        ("【界面呈现】", "资质校验弹窗通过；电子合同内嵌 e签宝 CA 数字证书认证，并显著提示‘30% 反跳单违约金条款’，定金 40% 冻结托管。"),
        ("【标准讲解话术】", "“飞手在小程序端抢单前，系统执行严格的执照类型、匹配机型和信用分校验。双方在线签署具备法律效力的电子合同，合同经过区块链存证并包含 30% 违约金条款，从源头建立契约保障。”")
    ])

    add_callout(doc, "步骤 2.3：履约合规 · UOM「一登二查三申请」与现场GPS打卡", [
        ("【演示操作】", "在 c-app.html 订单详情页中 ➔ 演示 UOM 步骤条（实名登记/空域查询/计划申报/飞行险）➔ 点击「现场 GPS 水印打卡」。"),
        ("【界面呈现】", "展示空域合规核验标记、阳光财产保险承保单号，GPS 坐标比对成功后点亮打卡状态。"),
        ("【标准讲解话术】", "“在作业起飞前，平台强制嵌入民航局 UOM 合规流程：核验飞手实名登记码与适飞空域属性，并附赠商业第三者责任险与机身险。飞手到达现场必须通过 GPS 水印打卡，杜绝虚假作业。”")
    ])

    add_callout(doc, "步骤 2.4：企业端调阅存证与微信分账结算", [
        ("【演示操作】", "切换至 b-enterprise.html ➔ 点击「电子合同与存证」➔ 点击「调阅电子存证证书」弹窗。"),
        ("【界面呈现】", "展示存证证书哈希防篡改码；验收后尾款结算至对公商户号，可申请 13% 增值税专票。"),
        ("【标准讲解话术】", "“作业完成后，企业在线调阅带可信时间戳的区块链存证凭证，确认成果后验收，微信支付分账系统自动完成佣金划扣与对公结算，平台支持开具增值税专票，完全契合大型企业财务制度。”")
    ])

    # 第三幕
    add_h2(doc, "第三幕：五维一体防跳单监控与仲裁法庭（技术护城河与风控中枢）")

    add_callout(doc, "步骤 3.1：C端 IM 沟通敏感词毫秒级语义拦截", [
        ("【演示操作】", "切换至 c-app.html ➔ 点击订单中的「IM沟通」➔ 演示飞手输入‘加我微信私下转账给你优惠’ ➔ 发送拦截。"),
        ("【界面呈现】", "系统立即弹出红色高危警示，文字自动转为‘***’脱敏，并提示‘已触发风控合规预警并上报运营中枢’。"),
        ("【标准讲解话术】", "“撮合平台最核心的商业风险是双方‘跳单私下交易’。毕数飞行平台自研了五维一体防跳单体系。首先在端侧，内置 IM 对手机号、微信号、私下结算等敏感词进行 NLP 毫秒级语义识别与实时拦截。”")
    ])

    add_callout(doc, "步骤 3.2：超管风控中枢 · 调阅隐私号通话录音与智能质检", [
        ("【演示操作】", "切换至 b-admin.html ➔ 点击左侧「防跳单监控与仲裁法庭」➔ 点击待处置案件「简阳 500 亩水稻植保跳单预警」的「调阅通话录音与智能质检」。"),
        ("【界面呈现】", "弹窗展示 AXB 隐私通话转文字明细，AI 自动标红‘走私单、少收平台费、微信转我’等确凿证据。"),
        ("【标准讲解话术】", "“平台所有双向通话均通过 AXB 隐私虚拟号中转。风控中枢通过语音转文字与智能质检算法，自动捕捉违规通话并将私单关键词标红提取，确凿固定证据链。”")
    ])

    add_callout(doc, "步骤 3.3：仲裁法庭处置 · 违约扣除 ¥1,000 保证金与信用制裁", [
        ("【演示操作】", "在录音质检弹窗底部点击「确认违规，执行仲裁处置」。"),
        ("【界面呈现】", "系统即时提示：扣除履约保证金 ¥1,000 作为违约金、扣除信用分 50 分、限制接单 30 天，并生成全网公示。"),
        ("【标准讲解话术】", "“风控专员一键下达仲裁决定：系统依据入驻协议，直接从飞手的微信商户保证金专户中划扣 ¥1,000 违约金，并降低其信用分，全平台公示，形成强大的法律与规则震慑。”")
    ])

    # 第四幕
    add_h2(doc, "第四幕：低空全域大屏与微信分账财务台账（商业模式与监管大盘）")

    add_callout(doc, "步骤 4.1：GIS 电子沙盘与实时空域态势大屏", [
        ("【演示操作】", "切换至 b-admin.html ➔ 点击左侧「低空运营大屏」。"),
        ("【界面呈现】", "展示全域 GIS 电子沙盘、实时空域态势图、订单与飞手分布、飞行架次/小时数及 UOM 越界违规监控。"),
        ("【标准讲解话术】", "“这是面向平台运营管理者与政府民航/公安监管部门定制的 GIS 电子沙盘大屏，全面实时掌控低空空域属性、任务执行热力与飞行安全态势。”")
    ])

    add_callout(doc, "步骤 4.2：维修服务商准入与无尘工位实拍核验", [
        ("【演示操作】", "点击左侧「准入审核与维修商」➔ 点击待审核服务商「成都翼修电子」的「资质核验」。"),
        ("【界面呈现】", "展示维修商资质证书、专用防静电无尘工位实拍图、技术团队认证与已冻结的履约保证金。"),
        ("【标准讲解话术】", "“在售后维修板块，我们对入驻维修商进行严格的实体工位实拍核验与保证金冻结，配合 C 端的 AI 故障预诊断与报价对比，为飞手提供标准化质保维修服务。”")
    ])

    add_callout(doc, "步骤 4.3：微信支付官方分账与财务清算台账（零银行、防二清）", [
        ("【演示操作】", "点击左侧「财务分账与清结算中枢」➔ 重点展示右上角‘微信分账冻结资金池 ¥2,450,000’ ➔ 点击「导出今日清结算总报表」。"),
        ("【界面呈现】", "展示 5 大板块阶梯抽佣明细流水（CAAC 20%、撮合 15%、维修 15%、保险 25%），全部走微信支付官方分账专线。"),
        ("【标准讲解话术】", "“最后是平台的资金合规与清算体系：平台全面采用微信支付官方分账（Profit Sharing）产品，待分账资金在微信商户号中自动冻结托管（最长30天），不经过任何第三方银行资金池，彻底杜绝‘二清’合规风险。平台每日自动按约定比例完成清分，支持一键导出财务对账单与 13% 专票抵扣明细。”")
    ])

    # 5. 答辩问答
    add_h1(doc, "四、 汇报常见问题与应答话术库（Q&A 备用）")
    
    qa_headers = ["序号", "领导 / 评委可能关注的问题", "核心应答要点与标准话术"]
    qa_rows = [
        ["1", "平台代收代付资金，会不会触碰央行‘二清’或非法集资红线？", 
         "【核心话术】：完全不会。平台已全面摒弃商业银行自建资金池模式，直接接入微信支付官方分账（Profit Sharing）产品。所有订单款项由微信持牌机构在商户号中原生冻结托管，验收后由微信官方接口直清分拨至接收方对公/个人账户，资金流向全程受微信支付持牌清算监管，天然合规。"],
        ["2", "飞手和客户如果私下打电话或微信转账（跳单），平台如何防范？", 
         "【核心话术】：我们构建了‘技术隔离+经济激励+信用约束+服务保障+法律震慑’五维一体防线：①C端IM聊天敏感词毫秒级NLP拦截；②双向通话全部强制走AXB隐私虚拟号并对录音进行AI质检；③电子合同内嵌30%违约金条款；④履约保证金违约直接扣除；⑤平台专属的飞行保险与电子存证仅对平台订单生效。"],
        ["3", "平台目前的开发完成度如何？下一步上线计划是什么？", 
         "【核心话术】：目前平台全套业务逻辑、4端高保真交互演示系统、C端 UniApp+Vue3 小程序前端源码工程及全套设计规范已 100% 准备就绪；后端 API 微服务架构与上线准备清单已全部定稿。下一步将并行推进微信商户号分账开通、人社定点合作确认与后端接口联调，预计 8-12 周即可正式上线运营。"],
        ["4", "平台与民航局 UOM 系统的关系是什么？是否替代政府审批？", 
         "【核心话术】：平台坚守‘合规引导者’定位，不替代 UOM 系统的空域管理与审批职能。平台在飞手接单与起飞前强制嵌入‘一登二查三申请’标准合规指引，核验登记码与适飞空域属性，留存合规凭证，协助政府主管部门实现低空飞行的规范化管理。"]
    ]
    add_table_data(doc, qa_headers, qa_rows, [1.2, 5.0, 10.3])

    # 保存文档
    doc.save(out_path)
    print(f"[OK] Word document successfully generated at: {out_path}")

if __name__ == '__main__':
    out_dir = r"d:\work\bishufei\设计文档"
    out_file = os.path.join(out_dir, "毕数飞行平台_系统演示汇报流程与讲解话术.docx")
    generate_demo_word_doc(out_file)
