# -*- coding: utf-8 -*-
"""
毕数飞行平台 - 小程序上线准备清单 Word 文档生成脚本
核心: 一张总表，三期事项全部放一起，让领导一目了然知道要准备什么
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ──────────────────────────────────────────────
# 辅助函数
# ──────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if level == 0:
            run.font.size = Pt(22)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        elif level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1A, 0x56, 0xC4)
        elif level == 2:
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return heading

def add_para(doc, text, bold=False, size=11, color=None, alignment=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run.font.size = Pt(size)
    return p

def add_table(doc, headers, rows, col_widths=None, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = '微软雅黑'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                run.font.size = Pt(10)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A56C4')

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(font_size)
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4FA')

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


# ──────────────────────────────────────────────
# 文档生成
# ──────────────────────────────────────────────

doc = Document()

style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

# ═══════════════════════════════════════════════
# 标题
# ═══════════════════════════════════════════════
add_heading(doc, '毕数飞行平台 - C端小程序上线准备清单', level=0)
add_para(doc, '编制日期: 2026年8月  |  用途: 领导层会议讨论, 安排人员申请和采购',
         size=11, color=RGBColor(0x66, 0x66, 0x66))
add_para(doc, '一期: CAAC考证报名撮合 + 飞手接单/用机需求撮合  |  '
              '二期: 电商商城 + 维修服务  |  '
              '三期: 技能培训补贴 + 风控深化',
         size=10, color=RGBColor(0x66, 0x66, 0x66))

# ═══════════════════════════════════════════════
# 一、全部需要准备的事项总表
# ═══════════════════════════════════════════════
add_heading(doc, '一、全部需要准备的事项总表', level=1)
add_para(doc, '下表列出三期全部需要申请/采购/准备的事项, 请领导统一安排人员推进。',
         bold=True, size=11, color=RGBColor(0xCC, 0x00, 0x00))

add_table(doc,
    ['期数', '类别', '事项', '所需材料/资质', '预计费用', '预计周期'],
    [
        # ── 账号与资质 ──
        ['一期', '账号资质', '注册微信小程序账号(企业主体), 获取 AppID', '营业执照、对公账户、法人信息', '免费', '1-5 工作日'],
        ['一期', '账号资质', '微信小程序认证', '营业执照、对公账户、法人信息', '300 元/年', '1-5 工作日'],
        ['一期', '账号资质', '小程序备案(工信部强制要求)', '营业执照、法人身份证、负责人信息、服务类目', '免费', '7-20 工作日'],
        ['一期', '账号资质', '域名购买 + ICP 备案', '营业执照', '约 100 元/年', '7-20 工作日'],
        ['一期', '账号资质', 'MP 后台选择服务类目: 商业服务-预约/上门服务', '营业执照', '免费', '1-3 工作日'],
        ['二期', '账号资质', '申请《增值电信业务经营许可证》(EDI 证)', '营业执照', '约 2000-5000 元(代办)', '20-40 工作日'],
        ['二期', '账号资质', 'MP 后台新增服务类目: 电商平台 + 维修/售后', '营业执照 + EDI 证', '免费', '1-3 工作日'],

        # ── 服务器与基础设施 ──
        ['一期', '基础设施', '采购云服务器(4核8G+)', '-', '约 300-800 元/月', '1-2 天'],
        ['一期', '基础设施', 'SSL 证书(HTTPS)', '-', '免费或 500+元/年', '1 天'],
        ['一期', '基础设施', 'MP 后台配置服务器域名', '域名已备案', '免费', '1 天'],

        # ── 微信支付与分账 ──
        ['一期', '支付分账', '申请微信支付商户号 + 开通 JSAPI 支付', '营业执照 + 对公账户', '免费', '1-3 工作日'],
        ['一期', '支付分账', '开通微信支付"分账"产品', '商户号就绪', '免费', '3-7 工作日'],
        ['一期', '支付分账', '添加分账接收方: 培训机构(企业对公)、飞手(个人)', '分账产品已开通', '免费', '1-3 天'],
        ['一期', '支付分账', '开通"商家转账到零钱"产品(飞手提现用)', '商户号就绪', '0.1% 手续费', '3-7 工作日'],
        ['二期', '支付分账', '添加分账接收方: 商家(企业对公)、维修商(企业对公)', '分账产品已开通', '免费', '1-3 天'],

        # ── 第三方服务采购 ──
        ['一期', '第三方服务', '云存储 OSS/COS(文件上传)', '-', '按量计费', '1 天'],
        ['一期', '第三方服务', '电子合同服务(e签宝/法大大)', '营业执照', '约 2-5 元/份', '3-5 天'],
        ['一期', '第三方服务', '虚拟号码服务(AXB 隐私通话)', '营业执照', '约 0.1 元/分钟', '3-5 天'],
        ['一期', '第三方服务', '保险服务(第三者责任险 + 机身险)', '营业执照', '按保单', '需与保司对接'],
        ['一期', '第三方服务', '短信服务(验证码/通知)', '营业执照', '约 0.04 元/条', '1-2 天'],
        ['一期', '第三方服务', '地图 SDK(腾讯位置服务)', '-', '免费', '1 天'],
        ['二期', '第三方服务', '物流对接 API(快递100/菜鸟面单)', '-', '按调用量计费', '2-3 天'],

        # ── 安全合规 ──
        ['一期', '安全合规', 'MP 后台填写隐私保护指引', '-', '免费', '1 天'],
        ['一期', '安全合规', '数据加密 + 前端脱敏 + 虚拟号码', '-', '免费', '开发实现'],
        ['三期', '安全认证', '等保三级认证(网络安全等级保护)', '安全体系就绪', '约 5-10 万元', '2-3 个月'],

        # ── 政府对接 ──
        ['三期', '政府对接', '人社部门补贴对接(确定落地城市 + 对接申报系统)', '需确认对接层级和资质', '-', '1-3 个月'],
        ['三期', '政府对接', 'UOM 平台数据直连(可选, 联系民航局)', '-', '-', '不确定'],
        ['三期', '政府对接', '公安报备系统对接(可选)', '-', '-', '不确定'],
    ],
    col_widths=[1.5, 2.5, 7, 5, 4, 4],
    font_size=8)

add_para(doc, '', size=4)
add_para(doc, '重要提示: EDI 证(20-40 天)和等保三级(2-3 个月)周期最长, '
              '建议在一期期间就并行启动申请, 不要等到二三期才开始办!',
         bold=True, size=11, color=RGBColor(0xCC, 0x00, 0x00))

# ═══════════════════════════════════════════════
# 二、分账场景一览
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '二、分账场景一览(全三期)', level=1)
add_para(doc, '平台共 6 条分账链路, 均通过微信支付分账产品实现。', bold=True, size=11,
         color=RGBColor(0xCC, 0x00, 0x00))

add_table(doc,
    ['期数', '#', '场景', '资金流向', '佣金比例'],
    [
        ['一期', '1', 'CAAC 考证报名', '学员支付 - 平台监管 - 平台抽佣 - 机构到账', '平台 20%, 机构 80%'],
        ['一期', '2', '订单撮合(定金+尾款)', '需求方支付 - 平台托管 - 验收后抽佣 - 飞手到账', '阶梯递减 10%-20%'],
        ['一期', '3', '飞手保证金管理', '飞手缴纳 - 平台资金池 - 违约扣除', '保证金沉淀, 非佣金'],
        ['二期', '4', '电商交易(商品)', '买家支付 - 平台 - 抽佣 - 卖家到账', '3%-8%'],
        ['二期', '5', '维修服务', '用户支付 - 平台托管 - 验收后抽佣 - 维修商到账', '5%-20%'],
        ['三期', '6', '技能培训补贴分成', '人社补贴 - 平台抽服务费 - 机构到账', '5%-10%'],
    ],
    col_widths=[2, 1, 4, 8, 4],
    font_size=9)

# ═══════════════════════════════════════════════
# 三、分账到账时间说明
# ═══════════════════════════════════════════════
add_heading(doc, '三、分账到账时间说明(重要)', level=1)
add_para(doc, '微信支付分账完成后资金是"准实时"到账, 但分账有两个前提条件:',
         bold=True, size=11, color=RGBColor(0xCC, 0x00, 0x00))
add_bullet(doc, '条件一: 订单需先结算。微信支付结算周期通常 T+1, 结算后才能调用分账接口')
add_bullet(doc, '条件二: 支付成功后需等待至少 1 分钟才能发起分账')
add_bullet(doc, '资金冻结期: 分账前资金冻结在商户号中, 最长 30 天, 超过 30 天自动解冻给平台')
add_bullet(doc, '分账完成后: 接收方可立即提现; 个人接收方分到零钱后可直接使用')

add_para(doc, '', size=4)
add_para(doc, '已修改的文案(原"实时到账"不准确, 已改为准确表述):', bold=True, size=11)
add_table(doc,
    ['页面', '原文案', '修改后'],
    [
        ['培训报名收银台', '平台服务费 20% 即结', '平台服务费 20% 次日结算'],
        ['培训报名收银台', '机构实时到账 80%', '机构 T+1 到账 80%'],
        ['CAAC 课程详情', '报名即分账 - 资金监管', '资金监管 - T+1 分账结算'],
        ['维修验收结算(二期)', '实时结算至维修商对公账户', 'T+1 结算至维修商对公账户'],
    ],
    col_widths=[4, 5.5, 6.5])

add_para(doc, '', size=4)
add_para(doc, '说明: 严格来说分账动作本身是准实时的, 但由于微信支付结算周期为 T+1, '
              '从用户支付到机构实际拿到钱, 整体周期是 T+1。文案修改为"T+1 到账"是合理的。',
         size=10, color=RGBColor(0x66, 0x66, 0x66))

# ═══════════════════════════════════════════════
# 四、资金监管与冻结机制
# ═══════════════════════════════════════════════
add_heading(doc, '四、资金监管与冻结机制', level=1)
add_para(doc, '微信支付分账产品自带资金冻结能力, 满足平台资金托管需求, 不需要单独开银行监管账户。')
add_bullet(doc, '资金冻结: 支付成功后, 待分账资金自动冻结在商户号中, 最长 30 天')
add_bullet(doc, '分账操作: 在冻结期内随时可发起分账(建议支付成功 1 分钟后)')
add_bullet(doc, '自动解冻: 超过 30 天未分账的订单, 资金自动解冻给分账方(平台)')
add_bullet(doc, '退款保障: 冻结期内退款直接从冻结资金出款, 不影响商户可用余额')

add_para(doc, '', size=4)
add_para(doc, '结论: 微信支付分账自带的冻结机制已能满足平台的资金托管需求(订单定金、培训费等)。', bold=True, size=11,
         color=RGBColor(0x1A, 0x56, 0xC4))

add_para(doc, '', size=4)
add_para(doc, '方案对比:', bold=True, size=11)
add_table(doc,
    ['方案', '优点', '缺点', '适用场景'],
    [
        ['微信支付分账(推荐)', '官方支持, 自带冻结, 开发量小, 免手续费', '单笔最高 30%, 结算后才能分账', '佣金分账 + 资金托管'],
        ['微信平台收付通', '自动分账到二级商户, 合规度高', '申请门槛高, 需电商平台资质', '二期电商板块'],
        ['商家转账到零钱', '灵活, 无比例限制', '非分账产品, 有额度限制, 有手续费', '飞手提现、保证金退还'],
    ],
    col_widths=[4, 5, 5, 5])

# ═══════════════════════════════════════════════
# 五、合规风险提示
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '五、合规风险提示(需领导知晓)', level=1)
add_table(doc,
    ['风险', '说明', '严重程度', '涉及期数'],
    [
        ['"二清"风险', '平台代收代付资金如果未取得支付牌照, 可能构成"无证经营支付业务"。'
         '使用微信支付分账产品可有效规避此风险, 资金冻结在微信商户号中, 由微信完成分账。', '高', '一期起'],
        ['资金池风险', '飞手保证金(500-2000 元/人)形成资金池, 大量飞手同时要求退还时'
         '需有足够流动资金。需设计保证金备付金机制。', '中', '一期起'],
        ['分账延迟风险', '微信支付分账 T+1 到账, UI 文案已修改为"T+1 到账", 避免用户投诉。', '中', '一期起'],
        ['退款资金链断裂', 'CAAC 退费如机构已分账到账, 平台需先垫付退款再向机构追讨。需设计退款准备金。', '中', '一期起'],
        ['发票与税务', '平台抽佣收入需开发票; 飞手个人收入涉及个税代扣代缴;'
         '企业客户需增值税专票。需财务团队确认税务流程。', '高', '一期起'],
        ['反洗钱', '大额分账、频繁提现可能触发反洗钱预警。需接入风控规则(单日限额、异常监测)。', '中', '一期起'],
        ['电商合规', 'EDI 证申请周期长(20-40 天); 二手商品交易可能需特殊资质; 7 天无理由退货法律适用。', '中', '二期'],
        ['保险经纪资质', '平台代收保费、代为投保, 可能涉及保险经纪业务。需确认是否需要保险经纪牌照。', '中', '二期'],
        ['人社补贴政策变动', '各地补贴政策不一且可能调整, "学员零费用"模式依赖补贴持续。需做政策风险预案。', '高', '三期'],
        ['等保三级周期长', '可能影响政府项目对接进度, 需尽早启动(2-3 个月)。', '中', '三期'],
    ],
    col_widths=[3, 12, 2, 2],
    font_size=8)

# ═══════════════════════════════════════════════
# 六、时限与关键路径
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '六、时限与关键路径', level=1)

add_heading(doc, '一期关键路径', level=2)
add_table(doc,
    ['事项', '前置条件', '预计耗时'],
    [
        ['域名注册 + ICP 备案', '营业执照', '7-20 个工作日'],
        ['小程序备案(工信部要求)', '小程序已注册 + 营业执照', '7-20 个工作日'],
        ['微信小程序注册 + 认证', '营业执照 + 对公账户', '1-5 个工作日'],
        ['微信支付商户号申请', '小程序已认证', '1-3 个工作日'],
        ['微信支付分账产品开通', '商户号就绪', '3-7 个工作日'],
        ['服务器采购 + HTTPS 配置', '域名备案完成', '1-2 天'],
        ['微信支付接入', '商户号就绪', '3-5 天'],
        ['真机测试 + 修复', '全部功能就绪', '1-2 周'],
        ['提交审核 + 发布', '以上全部完成', '1-7 天'],
    ],
    col_widths=[8, 6, 5])
add_para(doc, '', size=4)
add_para(doc, '一期总计最短周期: 约 8-12 周', bold=True, size=12,
         color=RGBColor(0xCC, 0x00, 0x00))

add_heading(doc, '二期关键路径(新增项)', level=2)
add_table(doc,
    ['事项', '前置条件', '预计耗时'],
    [
        ['EDI 证申请', '营业执照', '20-40 个工作日'],
        ['MP 后台新增电商类目', 'EDI 证就绪', '1-3 个工作日'],
        ['新增分账接收方(商家/维修商)', '分账产品已开通', '1-3 天'],
        ['物流 API 对接', '服务商签约', '2-3 天'],
        ['真机测试 + 审核发布', '以上完成', '1-2 周'],
    ],
    col_widths=[8, 6, 5])
add_para(doc, '', size=4)
add_para(doc, '二期新增周期: 约 6-10 周(EDI 证是关键瓶颈, 建议一期期间就启动申请)',
         bold=True, size=11, color=RGBColor(0xCC, 0x00, 0x00))

add_heading(doc, '三期关键路径(新增项)', level=2)
add_table(doc,
    ['事项', '前置条件', '预计耗时'],
    [
        ['人社部门对接', '确定落地城市', '1-3 个月(政府对接)'],
        ['等保三级测评', '安全体系就绪', '2-3 个月'],
        ['UOM 数据直连(可选)', '民航局沟通', '不确定(取决于政府)'],
        ['风控系统开发 + 测试', '后端团队就绪', '3-4 周'],
        ['真机测试 + 审核发布', '以上完成', '1-2 周'],
    ],
    col_widths=[8, 6, 5])
add_para(doc, '', size=4)
add_para(doc, '三期新增周期: 约 3-5 个月(政府对接是最大不确定项)',
         bold=True, size=11, color=RGBColor(0xCC, 0x00, 0x00))

add_para(doc, '', size=6)
add_para(doc, '建议: EDI 证和等保三级周期最长, 建议在一期期间就并行启动申请, '
              '不要等到二三期才开始办。', bold=True, size=11,
         color=RGBColor(0xCC, 0x00, 0x00))

# ═══════════════════════════════════════════════
# 七、保证金与退款说明
# ═══════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, '七、保证金与退款说明', level=1)

add_heading(doc, '7.1 飞手保证金体系(一期)', level=2)
add_para(doc, '飞手保证金 500-2000 元/人。')
add_bullet(doc, '飞手缴纳保证金: 通过微信支付充值到平台商户号')
add_bullet(doc, '保证金冻结: 后端标记为"已冻结", 不可提现')
add_bullet(doc, '违约扣除: 跳单核实后扣除, 需有仲裁记录和用户确认')
add_bullet(doc, '保证金退还: 飞手退出平台时全额退还')

add_heading(doc, '7.2 飞手提现(一期)', level=2)
add_bullet(doc, '使用微信支付"商家转账到零钱"产品(V3 API)')
add_bullet(doc, '注意: 微信转账单笔上限 20000 元, 日上限 200000 元, 有 0.1% 手续费')

add_heading(doc, '7.3 退款场景', level=2)
add_table(doc,
    ['期数', '场景', '说明'],
    [
        ['一期', 'CAAC 考试未通过(阶梯退费)', '部分退款, 需后端支持阶梯退费逻辑'],
        ['一期', '订单取消/纠纷仲裁退定金', '全额退款, 从冻结资金出款'],
        ['二期', '电商 7 天无理由退货', '全额退款'],
        ['二期', '维修纠纷退托管款', '从冻结资金出款'],
    ],
    col_widths=[2, 6, 9])
add_para(doc, '需要准备: 后端实现微信支付退款接口(V3 API), 支持全额和部分退款, 处理退款回调通知。', size=10,
         color=RGBColor(0x66, 0x66, 0x66))

# ──────────────────────────────────────────────
# 保存
# ──────────────────────────────────────────────
output_path = r'd:\work\bishufei\设计文档\毕数飞行平台_小程序上线准备清单_v10.docx'
doc.save(output_path)
print(f'Word doc generated: {output_path}')
